from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from pathlib import Path

out = Path(r"D:\Portfolio2026\resume-output\Daniel_Chavez_Cover_Letter_Ramp.docx")
doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.8)
sec.left_margin = Inches(0.9)
sec.right_margin = Inches(0.9)

normal = doc.styles["Normal"]
normal.font.name = "Aptos"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
normal.font.size = Pt(10.8)
normal.font.color.rgb = RGBColor.from_string("222222")
normal.paragraph_format.space_after = Pt(9)
normal.paragraph_format.line_spacing = 1.12

def add(text, size=10.8, color="222222", bold=False, after=9):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    r.font.name = "Aptos"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    r.bold = bold
    return p

footer = sec.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run("Daniel Chavez  |  Cover Letter  |  Ramp")
fr.font.size = Pt(8.5)
fr.font.color.rgb = RGBColor.from_string("718096")

add("DANIEL CHAVEZ", 22, "173B57", True, 2)
add("Winnipeg, MB  |  204-721-2248  |  dchavez3395@gmail.com", 9.5, "5C6B73", False, 18)
add("July 14, 2026", 10.8, after=14)
add("Hiring Team", 10.8, "222222", True, 1)
add("Ramp", 10.8, after=1)
add("Re: Software Engineer, Frontend", 10.8, "173B57", True, 14)

add("Dear Ramp hiring team,", 10.8, "222222", True, 10)
add("I am excited to apply for the Software Engineer, Frontend position. I am a Winnipeg-based web developer with professional agency and freelance experience building responsive, accessible, and performance-focused web experiences with React, JavaScript, HTML, CSS/SCSS, REST APIs, Git/GitHub, and modern deployment workflows.")
add("My work has involved turning product and design requirements into polished interfaces, reusable frontend components, reliable integrations, and production-ready websites. I care about getting the details right: visual fidelity, keyboard accessibility, responsive behavior, performance, cross-browser quality, and clear communication with designers, product partners, and other developers. My portfolio at danieldev.ca reflects that hands-on approach, including work deployed through Vercel and maintained through Git-based workflows.")
add("Ramp's focus on building high-quality software that helps businesses operate more efficiently is especially appealing to me. I would bring strong frontend fundamentals, practical product judgment, agency experience, and a habit of taking ownership from implementation through testing and refinement. I am comfortable learning an existing design system and codebase quickly, and I would be excited to contribute to a React-based frontend stack while continuing to deepen my TypeScript and product-engineering experience.")
add("Thank you for your consideration. I would welcome the opportunity to discuss how my experience could contribute to Ramp's frontend engineering team.", 10.8, after=12)
add("Sincerely,", 10.8, after=2)
add("Daniel Chavez", 10.8, "173B57", True, 1)
add("Portfolio: https://danieldev.ca   |   GitHub: https://github.com/dchavez3395", 9.5, "5C6B73", False, 1)
add("LinkedIn: https://www.linkedin.com/in/danielchavez2/", 9.5, "5C6B73", False, 0)

doc.core_properties.title = "Cover Letter - Ramp Software Engineer, Frontend"
doc.core_properties.subject = "Daniel Chavez application for Ramp"
doc.core_properties.author = "Daniel Chavez"
out.parent.mkdir(parents=True, exist_ok=True)
doc.save(out)
print(out)
