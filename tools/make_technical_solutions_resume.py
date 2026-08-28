from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "resume-output"
OUT.mkdir(exist_ok=True)
FINAL_DOCX = OUT / "Daniel_Chavez_Technical_Solutions_Resume.docx"

# compact_reference_guide with a named resume_density override:
# Letter portrait; 0.65-inch margins; Calibri 9.4 pt body; 1.05 spacing.
# The override keeps a career document to two pages without shrinking below
# normal resume typography. Headings and real list numbering remain explicit.
INK = RGBColor(27, 31, 36)
BLUE = RGBColor(20, 83, 111)
MUTED = RGBColor(88, 96, 105)
HAIRLINE = "D7E1E7"
BODY_SIZE = 9.4
CONTENT_WIDTH = 7.2


def set_font(run, size=BODY_SIZE, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:ascii"), "Calibri")
    rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_spacing(paragraph, before=0, after=3, line=1.05, keep_next=False):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.keep_with_next = keep_next


def add_paragraph(doc, text, size=BODY_SIZE, bold=False, color=INK, before=0,
                  after=3, line=1.05, italic=False, align=None):
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    set_spacing(paragraph, before, after, line)
    set_font(paragraph.add_run(text), size, bold, color, italic)
    return paragraph


def add_numbering_definition(doc):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=-1) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "-")
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "540")
    indent.set(qn("w:hanging"), "270")
    ppr.append(tabs)
    ppr.append(indent)
    for child in (start, num_fmt, level_text, suffix, ppr):
        level.append(child)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullet(doc, num_id, text, after=1.8):
    paragraph = doc.add_paragraph()
    set_spacing(paragraph, 0, after, 1.03)
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num)
    ppr.append(num_pr)
    set_font(paragraph.add_run(text), BODY_SIZE)
    return paragraph


def add_section_heading(doc, text):
    paragraph = doc.add_paragraph(style="Resume Section")
    set_spacing(paragraph, 7, 3, 1.0, keep_next=True)
    set_font(paragraph.add_run(text.upper()), 10.2, True, BLUE)
    return paragraph


def add_job(doc, num_id, role, company, location, dates, bullets):
    heading = doc.add_paragraph()
    set_spacing(heading, 4, 0, 1.0, keep_next=True)
    heading.paragraph_format.tab_stops.add_tab_stop(
        Inches(CONTENT_WIDTH), WD_TAB_ALIGNMENT.RIGHT
    )
    set_font(heading.add_run(role), 10.1, True)
    set_font(heading.add_run(f" - {company}"), 10.1, True, BLUE)
    set_font(heading.add_run(f"\t{dates}"), 8.8, False, MUTED)

    meta = doc.add_paragraph()
    set_spacing(meta, 0, 1.5, 1.0, keep_next=True)
    set_font(meta.add_run(location), 8.8, False, MUTED, italic=True)
    for bullet in bullets:
        add_bullet(doc, num_id, bullet)


def add_capability(doc, label, text):
    paragraph = doc.add_paragraph()
    set_spacing(paragraph, 0, 2, 1.03)
    set_font(paragraph.add_run(f"{label}: "), 9.2, True)
    set_font(paragraph.add_run(text), 9.2)


def build():
    doc = Document()
    doc.core_properties.title = "Daniel Chavez - Technical Solutions Resume"
    doc.core_properties.subject = "Sales engineering, solutions, implementation, and client delivery"
    doc.core_properties.author = "Daniel Chavez"

    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.05

    if "Resume Section" not in [style.name for style in doc.styles]:
        section_style = doc.styles.add_style("Resume Section", WD_STYLE_TYPE.PARAGRAPH)
    else:
        section_style = doc.styles["Resume Section"]
    section_style.font.name = "Calibri"
    section_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    section_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    section_style.font.size = Pt(10.2)
    section_style.font.bold = True
    section_style.font.color.rgb = BLUE

    num_id = add_numbering_definition(doc)

    # customer_pack-inspired candidate masthead: left-aligned, commercial, compact.
    add_paragraph(doc, "DANIEL CHAVEZ", 23, True, INK, after=0, line=1.0)
    add_paragraph(
        doc,
        "TECHNICAL SOLUTIONS | CLIENT DELIVERY | SALES",
        11.2,
        True,
        BLUE,
        after=2,
        line=1.0,
    )
    add_paragraph(
        doc,
        "Winnipeg, MB | 204-721-2248 | dchavez3395@gmail.com | danieldev.ca | linkedin.com/in/danielchavez2",
        8.9,
        False,
        MUTED,
        after=6,
        line=1.0,
    )

    add_section_heading(doc, "Professional Summary")
    add_paragraph(
        doc,
        "Bilingual technical-solutions professional combining 4+ years of production web delivery with automotive sales, finance, CRM, prospecting, and relationship ownership. Experienced translating customer and stakeholder needs into accessible web applications, CMS workflows, API integrations, and dependable implementation. Strong fit for sales engineering, solutions consulting, implementation, technical account management, and customer-facing SaaS roles.",
        9.5,
        after=3,
        line=1.06,
    )

    add_section_heading(doc, "Core Capabilities")
    add_capability(doc, "Customer and revenue", "Consultative discovery, CRM and pipeline ownership, product explanation, objection handling, negotiation, follow-up, automotive F&I")
    add_capability(doc, "Solutions and delivery", "Requirements translation, client walkthroughs, REST APIs, CMS workflows, implementation, troubleshooting, stakeholder updates, QA, training and support")
    add_capability(doc, "Technical foundation", "JavaScript, React, Node.js, Express, PHP, WordPress, Shopify Hydrogen, HTML/CSS, Git/GitHub, Figma, Vercel")
    add_capability(doc, "Quality and risk", "WCAG accessibility, semantic HTML, documentation, compliance mindset, SEO, performance, cross-browser testing, AI-assisted workflows")

    add_section_heading(doc, "Professional Experience")
    add_job(doc, num_id, "Web Developer", "Vincent Design Inc.", "Winnipeg, MB", "Jan 2023 - Jul 2026", [
        "Built and maintained responsive production websites and web applications for agency clients across annual reports, Indigenous organizations, non-profits, public information, aviation, and e-commerce.",
        "Translated Figma designs and stakeholder requirements into reusable React and JavaScript components, WordPress/CMS workflows, REST API integrations, and scalable SCSS and Tailwind systems.",
        "Supported projects from build through launch and ongoing improvement, including accessibility QA, troubleshooting, deployment, content workflows, SEO, and performance across Vercel, WP Engine, and Flywheel.",
    ])
    add_job(doc, num_id, "Front-End Developer", "Freelance", "Winnipeg, MB", "Aug 2021 - Present", [
        "Turn design mockups and existing codebases into responsive applications; build reusable components, add AJAX/JSON functionality, and debug or refactor legacy front-end code.",
    ])
    add_job(doc, num_id, "Special Finance Manager", "Birchwood Credit Solutions", "Winnipeg, MB", "Nov 2018 - Jul 2021", [
        "Structured automotive finance deals across prime, near-prime, and subprime credit tiers, balancing customer affordability, lender requirements, documentation, compliance, and dealership objectives.",
        "Managed the process from credit submission through funding using Dealertrack, Equifax reports, lender negotiation, stipulation management, menu presentation, and contract completion.",
    ])
    add_job(doc, num_id, "Business Development Manager", "Birchwood Nissan Brandon", "Brandon, MB", "Jan 2018 - Nov 2018", [
        "Managed inbound and outbound sales pipelines, CRM lead workflows, appointment conversion, and follow-up discipline while coordinating with sales and finance teams.",
    ])
    add_job(doc, num_id, "Sales Consultant", "Birchwood Nissan Brandon", "Brandon, MB", "May 2015 - Jan 2018", [
        "Used consultative discovery, product knowledge, negotiation, and disciplined follow-up to guide customers through high-consideration vehicle purchases.",
        "Maintained the dealership's highest closing ratio in 2016 and 2018 and received the Bronze Nissan Club Excellence Award in 2016.",
    ])

    add_section_heading(doc, "Additional Sales Background")
    add_paragraph(
        doc,
        "Combined Insurance - door-to-door insurance prospecting and sales | House of Nissan - automotive sales | The Source - consumer technology retail sales",
        9.3,
        after=3,
        line=1.04,
    )

    add_section_heading(doc, "Selected Technical Work")
    add_bullet(doc, num_id, "Puchica - Built a Shopify Hydrogen storefront from an empty repository with reusable React components, catalogue taxonomy, SEO infrastructure, and Shopify Admin API integration.")
    add_bullet(doc, num_id, "TELUS Reconciliation Report - Built an accessible React annual-report experience with screen-reader-aware data visualizations, client-controlled publishing, and an ongoing CMS workflow.")
    add_bullet(doc, num_id, "Tamalpais Trust - Delivered a WordPress grant-partner directory, interactive map, community-investments showcase, and editor-run publishing workflow.")

    add_section_heading(doc, "Training and Languages")
    add_bullet(doc, num_id, "Web Accessibility Training - Deque University")
    add_bullet(doc, num_id, "Complete Full-Stack Web Development Bootcamp - Dr. Angela Yu, Udemy")
    add_bullet(doc, num_id, "Modern PHP: Beginner to Advanced - Jannis Seemann, Udemy")
    add_bullet(doc, num_id, "English and Spanish - fluent professional communication")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(
        footer.add_run("Daniel Chavez | Technical Solutions | danieldev.ca"),
        8,
        False,
        MUTED,
    )

    doc.save(FINAL_DOCX)
    print(FINAL_DOCX)


if __name__ == "__main__":
    build()
