from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "resume-output"
OUT.mkdir(exist_ok=True)

BLUE = RGBColor(20, 83, 111)
INK = RGBColor(35, 35, 35)
MUTED = RGBColor(95, 95, 95)

def font(run, size=10, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color

def pformat(p, before=0, after=4, line=1.05):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line

def add_text(doc, text, size=10, bold=False, color=INK, align=None, before=0, after=4, line=1.05, italic=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pformat(p, before, after, line)
    r = p.add_run(text)
    font(r, size, bold, color, italic)
    return p

def add_heading(doc, text):
    p = doc.add_paragraph()
    pformat(p, 8, 3, 1.0)
    r = p.add_run(text.upper())
    font(r, 10, True, BLUE)
    pPr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D6E2E8")
    borders.append(bottom)
    pPr.append(borders)

def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    pformat(p, 0, 2, 1.03)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    r = p.add_run(text)
    font(r, 9.5, False, INK)

def add_job(doc, role, company, dates, bullets):
    p = doc.add_paragraph()
    pformat(p, 5, 1, 1.0)
    r = p.add_run(role + " — ")
    font(r, 10, True, INK)
    r = p.add_run(company)
    font(r, 10, True, BLUE)
    r = p.add_run("  |  " + dates)
    font(r, 9, False, MUTED)
    for b in bullets:
        add_bullet(doc, b)

def build(with_references):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.55)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.3)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    styles["Normal"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    styles["Normal"].font.size = Pt(9.5)
    styles["List Bullet"].font.name = "Calibri"
    styles["List Bullet"]._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    styles["List Bullet"]._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    styles["List Bullet"].font.size = Pt(9.5)

    add_text(doc, "Daniel Chavez", 23, True, INK, WD_ALIGN_PARAGRAPH.CENTER, 0, 1, 1.0)
    add_text(doc, "Web Developer  |  Front-End / React / CMS", 11, True, BLUE, WD_ALIGN_PARAGRAPH.CENTER, 0, 2, 1.0)
    add_text(doc, "Winnipeg, MB  •  204-721-2248  •  dchavez3395@gmail.com  •  danieldev.ca", 9, False, MUTED, WD_ALIGN_PARAGRAPH.CENTER, 0, 7, 1.0)

    add_heading(doc, "Professional Summary")
    add_text(doc, "Web developer with 4+ years of professional and freelance experience building and maintaining production websites and web applications. Strongest in JavaScript, React, Node.js, HTML, CSS/SCSS, WordPress, and CMS-driven delivery, with practical experience across design handoff, accessibility, QA, APIs, SEO, performance, and deployment. Based in Winnipeg and seeking junior-to-mid front-end, web developer, React, full-stack web, CMS, or accessibility-focused roles.", 9.5, False, INK, None, 0, 4, 1.08)

    add_heading(doc, "Core Skills")
    add_text(doc, "JavaScript (ES6+), React, Node.js, Express, REST APIs, HTML5, CSS, SCSS/SASS, Tailwind CSS, PHP, WordPress custom themes, EJS, Blade, Git/GitHub, Figma handoff, Vercel, WP Engine, Flywheel, WCAG accessibility, semantic HTML, keyboard QA, SEO, performance optimization, cross-browser testing, Cursor, GitHub Copilot, Claude", 9.2, False, INK, None, 0, 4, 1.05)

    add_heading(doc, "Work Experience")
    add_job(doc, "Web Developer", "Vincent Design Inc.", "Winnipeg, MB  |  Jan 2023 – Jul 2026", [
        "Built and maintained responsive production websites and web applications for agency clients across annual reports, Indigenous organizations, non-profits, public information, and e-commerce.",
        "Developed reusable React and JavaScript components, scalable SCSS and Tailwind styling systems, and data-driven interfaces connected to REST APIs and CMS platforms.",
        "Partnered directly with designers in Figma to deliver semantic, accessible, responsive, and performant interfaces using WCAG as a working benchmark.",
        "Owned deployments and hosting through Vercel, WP Engine, and Flywheel, including troubleshooting, content workflows, SEO, and performance improvements.",
    ])
    add_job(doc, "Front-End Developer", "Freelance", "Winnipeg, MB  |  Aug 2021 – Present", [
        "Convert design mockups into responsive, cross-browser applications with React, JavaScript, HTML, and CSS.",
        "Build reusable component libraries and interactive functionality with AJAX and JSON; debug, refactor, and optimize legacy codebases.",
    ])
    add_job(doc, "Special Finance Manager", "Birchwood Credit Solutions", "Winnipeg, MB  |  Nov 2018 – Jul 2021", [
        "Structured automotive finance deals across all credit tiers and managed documentation from submission to funding in detail-heavy digital systems.",
    ])
    add_job(doc, "Business Development Manager", "Birchwood Nissan Brandon", "Brandon, MB  |  Jan 2018 – Nov 2018", [
        "Managed sales pipelines and CRM lead workflows, strengthening appointment conversion and follow-up discipline.",
    ])

    add_heading(doc, "Selected Projects")
    add_bullet(doc, "DanielDev Portfolio — React, Vite, Tailwind CSS, Vercel. Personal portfolio presenting production project work, accessibility notes, and case studies.")
    add_bullet(doc, "Puchica — Shopify Hydrogen, React. Built a custom storefront from an empty repository with reusable components, collection taxonomy, SEO infrastructure, and Shopify Admin API integration.")
    add_bullet(doc, "TELUS Reconciliation Report — React, headless CMS. Built an accessible annual-report experience with client-controlled publishing and data visualizations.")
    add_bullet(doc, "Tamalpais Trust — WordPress, interactive map. Built a grant-partner directory, interactive map, community-investments showcase, and editor-run publishing workflow.")

    add_heading(doc, "Certifications & Training")
    add_bullet(doc, "Complete Full-Stack Web Development Bootcamp — Dr. Angela Yu, Udemy")
    add_bullet(doc, "Modern PHP: Beginner to Advanced — Jannis Seemann, Udemy")
    add_bullet(doc, "Web Accessibility Training — Deque University")
    add_bullet(doc, "English and Spanish — Fluent (C2)")

    if with_references:
        add_heading(doc, "Professional References")
        refs = [
            ("Keith Solomon", "Senior Developer", "204-298-9732", "Keith@keithsolomon.net"),
            ("Cierra Waddell-Hodgson", "General Sales Manager", "204-573-7460", "cierra.waddell@hotmail.com"),
            ("Brady Sobering", "Sales Manager", "204-612-2497", "brady.sobering@gmail.com"),
        ]
        for name, title, phone, email in refs:
            p = doc.add_paragraph()
            pformat(p, 3, 2, 1.0)
            r = p.add_run(name + " — ")
            font(r, 9.5, True, INK)
            r = p.add_run(title + "  |  " + phone + "  |  " + email)
            font(r, 9.5, False, MUTED)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Daniel Chavez  •  Web Developer  •  danieldev.ca")
    font(r, 8, False, MUTED)

    name = "Daniel_Chavez_Resume_with_References.docx" if with_references else "Daniel_Chavez_Resume_Public.docx"
    path = OUT / name
    doc.save(path)
    print(path)

build(False)
build(True)
